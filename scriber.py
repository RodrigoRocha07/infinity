from docx import Document
import os

lista_clientes = [
{
    'nome':'Rodrigo Flávio Nunes Rocha',
    'cpf':'123.123.123-21',
    'rg':'123456789789',
    'endereco':' Rua Uruburetama 390',
    'cep':'60410306',
    'telefone':'(85) 99417-2862',
    'email':'rodrigorochafn@gmail.com'
},
{
    'nome':'Ana Elisa',
    'cpf':'321.321.321-21',
    'rg':'99999999999',
    'endereco':' Rua Uruburetama 390',
    'cep':'60410306',
    'telefone':'(85) 99417-2862',
    'email':'anaelisa@gmail.com'
}
]

pasta = 'contratos'
os.makedirs(pasta, exist_ok=True)

for cliente in lista_clientes:
    doc = Document()
    doc.add_heading('CONTRATO DE PRESTAÇÃO DE SERVIÇOS EM SAÚDE', 0)
    primeiro_nome = cliente['nome'].split()[0]
    doc.add_paragraph('I – DAS PARTES ')
    doc.add_paragraph('PACIENTE (PACIENTE):')
    doc.add_paragraph(f'Nome: {cliente["nome"]}')
    doc.add_paragraph(f'CPF: {cliente["cpf"]}')
    doc.add_paragraph(f'RG: {cliente["rg"]}')
    doc.add_paragraph(f'Endereco : {cliente["endereco"]}')
    doc.add_paragraph(f'Cep : {cliente["cep"]}')
    doc.add_paragraph(f'Telefone : {cliente["telefone"]}')
    doc.add_paragraph(f'Email : {cliente["email"]}')

    caminho_arquivo = os.path.join(pasta,f'contrato{primeiro_nome}.docx')
    doc.save(caminho_arquivo)






